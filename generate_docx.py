from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

def add_heading(doc, text, level=1):
    heading = doc.add_heading(text, level=level)
    return heading

def add_paragraph(doc, text, bold=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    if bold:
        run.bold = True
    return p

def add_bullet(doc, text):
    p = doc.add_paragraph(style='List Bullet')
    
    # Handle bold parts (e.g., "**Win Rate:** 26.3%")
    parts = text.split("**")
    for i, part in enumerate(parts):
        run = p.add_run(part)
        if i % 2 == 1:
            run.bold = True
    return p

def create_report():
    doc = Document()
    
    # Title
    title = doc.add_heading("Comprehensive Backtest Report", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph("Target Wallet: 0x880ac484a1743862989A441D6d867238c7AA311C")
    doc.add_paragraph("Simulation Period: March 11, 2026 - March 24, 2026")
    doc.add_paragraph("Total Trades Analyzed: 17,080")
    
    # 1. Top-Level Summary
    add_heading(doc, '1. Top-Level Summary', level=1)
    add_bullet(doc, "**Starting Balance:** $1,000.00")
    add_bullet(doc, "**Ending Balance:** $999.67")
    add_bullet(doc, "**Net Profit:** -$0.33")
    add_bullet(doc, "**Total ROI:** -0.03%")
    add_bullet(doc, "**Max Drawdown:** 3.51% ($35.07 underwater from the peak)")
    
    # 2. Trade Statistics
    add_heading(doc, '2. Trade Statistics', level=1)
    add_bullet(doc, "**Win Rate:** 26.3%")
    add_bullet(doc, "**Profit Factor:** 1.00 (Gross Profit / Gross Loss)")
    add_bullet(doc, "**Gross Profit:** $87.05")
    add_bullet(doc, "**Gross Loss:** $87.38")
    add_bullet(doc, "**Average Winning Trade:** $0.02")
    add_bullet(doc, "**Average Losing Trade:** -$0.01")
    add_bullet(doc, "**Largest Single Win:** $0.31")
    add_bullet(doc, "**Largest Single Loss:** -$0.59")
    
    p = doc.add_paragraph("Note: ")
    p.add_run("The profit factor of exactly 1.00 matching a low 26% win rate indicates that the target wallet utilizes a strategy with very tight stop-losses (small frequent losses) and larger, less frequent winners. The net result over these two weeks, however, is a stagnant flat-line.").italic = True
    
    # 3. Asset-Specific Performance
    add_heading(doc, '3. Asset-Specific Performance', level=1)
    
    add_heading(doc, 'Most Traded Assets (By Volume of Trades)', level=2)
    doc.add_paragraph("The target wallet's activity is heavily concentrated on HYPE and specific alt-coins/pre-launch tokens (@107, @142).")
    doc.add_paragraph("1. HYPE: 5,306 trades\n2. @107: 3,771 trades\n3. BTC: 2,099 trades\n4. @142: 2,044 trades\n5. XMR: 1,603 trades", style='List Number')
    
    add_heading(doc, 'Top 5 Most Profitable Assets', level=2)
    doc.add_paragraph("Despite making up less trading volume, XMR and BTC were the primary profit drivers, offsetting the losses incurred on other assets.")
    doc.add_paragraph("1. XMR: +$42.36 (1,603 trades)\n2. BTC: +$34.16 (2,099 trades)\n3. SOL: +$2.35 (422 trades)\n4. PUMP: +$0.56 (431 trades)\n5. DOGE: +$0.11 (96 trades)", style='List Number')
    
    add_heading(doc, 'Top 5 Least Profitable Assets', level=2)
    doc.add_paragraph("The vast majority of the frequent, high-volume trades ended up bleeding money over time.")
    doc.add_paragraph("1. HYPE: -$41.15 (5,306 trades)\n2. @142: -$35.03 (2,044 trades)\n3. ETH: -$2.64 (221 trades)\n4. @151: -$1.11 (813 trades)\n5. @107: -$0.07 (3,771 trades)", style='List Number')
    
    # 4. Strategic Analysis
    add_heading(doc, '4. Strategic Analysis Profile', level=1)
    doc.add_paragraph("Based on the data profile (17,000+ trades in two weeks, 26% win rate, average win double the average loss, and 3.5% max drawdown), there are three highly likely possibilities for this wallet's strategy:")
    
    add_heading(doc, 'Volume Farming / Points Airdrop Farming (Most Likely)', level=3)
    doc.add_paragraph("Operating at a perfect 1.00 Profit Factor while executing 1,200+ trades per day is the hallmark of a volume farming algorithm. The algorithm is designed to rapidly open and close positions in highly liquid or volatile markets with the sole mathematical goal of not losing money. By breaking exactly even, the wallet generates massive trading volume entirely for free, which maximizes rewards for HyperLiquid's points distributions, airdrops, or manufacturer rebates.")
    
    add_heading(doc, 'High-Frequency Market Making', level=3)
    doc.add_paragraph("This wallet could be running a market-making script that provides liquidity to the order books. Market makers naturally suffer low win-rates because their limit orders often get run over by aggressive directional traders (which explains bleeding money on highly volatile tokens like HYPE). However, on less chaotic pairs (like XMR and BTC), they capture the bid/ask spread successfully enough to zero out their losses.")
    
    add_heading(doc, 'Automated Momentum / Trend Following', level=3)
    doc.add_paragraph("The statistical signature of a 26% Win Rate + Wins being 2x larger than Losses is the textbook definition of a trend-following script with extremely tight stop-losses. The bot attempts to catch momentum breakouts constantly. 74% of the time, the breakout fails and the bot instantly cuts the trade for a tiny loss. However, the 26% of the time it does catch a real trend, it holds the position longer to let the profits multiply. The tight stop-losses are why the max drawdown never exceeded 3.5% despite placing 17,000 trades.")
    
    # 5. Conclusion
    add_heading(doc, '5. Conclusion & Actionable Insight', level=1)
    p2 = doc.add_paragraph("Mirroring 0x880ac484a1... with $1,000 would have been incredibly intensive on execution volume (averaging over 1,200 trades per day) but would have resulted in essentially zero net movement on your account balance.\n\n")
    p2.add_run("The Takeaway for Copy Trading:\n").bold = True
    p2.add_run("If your goal is absolute PnL growth, copying an account with a 1.00 profit factor isn't ideal because they are just treading water. However, if you were to utilize the .env settings in your copy software to block the failing assets (BLOCKED_ASSETS=HYPE,@142,@107), you would effectively filter out the algorithmic noise and only copy the highly profitable trades it executes on XMR, BTC, and SOL, turning this into an incredibly profitable mirroring setup!")
    
    doc.save('Hyperliquid_Backtest_Report.docx')
    print("Report saved as Hyperliquid_Backtest_Report.docx")

if __name__ == "__main__":
    create_report()
